

from fastapi import FastAPI, File, UploadFile, Form, Request, HTTPException
from fastapi.responses import HTMLResponse, JSONResponse
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles
from transformers import pipeline
import PyPDF2
import docx
import io
from typing import Optional

app = FastAPI()
templates = Jinja2Templates(directory="templates")

# Load your fine-tuned Hugging Face model
try:
    scorer = pipeline("text-classification", model="riddhii29/results")
except Exception as e:
    print(f"Warning: Could not load model. Using demo mode. Error: {e}")
    scorer = None

def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file"""
    try:
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading PDF: {str(e)}")

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        doc_file = io.BytesIO(file_content)
        doc = docx.Document(doc_file)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text.strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading DOCX: {str(e)}")

def extract_text_from_txt(file_content: bytes) -> str:
    """Extract text from TXT file"""
    try:
        return file_content.decode('utf-8').strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading TXT: {str(e)}")

async def process_file(file: UploadFile) -> str:
    """Process uploaded file and extract text"""
    if not file:
        raise HTTPException(status_code=400, detail="No file provided")
    
    content = await file.read()
    filename = file.filename.lower()
    
    if filename.endswith('.pdf'):
        return extract_text_from_pdf(content)
    elif filename.endswith('.docx'):
        return extract_text_from_docx(content)
    elif filename.endswith('.txt'):
        return extract_text_from_txt(content)
    else:
        raise HTTPException(
            status_code=400, 
            detail="Unsupported file format. Please upload PDF, DOCX, or TXT file."
        )

def calculate_fit_score(resume_text: str, job_description: str) -> dict:
    """Calculate fit score using the model or demo scoring"""
    if scorer:
        # Truncate text to fit model's max length (512 tokens ≈ 2000 characters)
        max_chars_per_section = 1000  # Leave room for [SEP] token
        
        resume_truncated = resume_text[:max_chars_per_section]
        job_truncated = job_description[:max_chars_per_section]
        
        input_text = f"{resume_truncated} [SEP] {job_truncated}"
        
        # Truncate to max_length with the tokenizer
        result = scorer(input_text, truncation=True, max_length=512)[0]
        fit_score = round(float(result['score']) * 10, 1)
        confidence = round(float(result['score']) * 100, 1)
    else:
        # Demo scoring based on keyword matching
        resume_lower = resume_text.lower()
        job_lower = job_description.lower()
        
        job_keywords = set(job_lower.split())
        resume_keywords = set(resume_lower.split())
        
        common_keywords = job_keywords.intersection(resume_keywords)
        match_ratio = len(common_keywords) / max(len(job_keywords), 1)
        
        fit_score = round(min(match_ratio * 10, 10), 1)
        confidence = round(match_ratio * 100, 1)
    
    # Determine rating category
    if fit_score >= 8:
        rating = "Excellent Match"
        color = "#10b981"
    elif fit_score >= 6:
        rating = "Good Match"
        color = "#3b82f6"
    elif fit_score >= 4:
        rating = "Fair Match"
        color = "#f59e0b"
    else:
        rating = "Poor Match"
        color = "#ef4444"
    
    return {
        "fit_score": fit_score,
        "confidence": confidence,
        "rating": rating,
        "color": color
    }

@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})

@app.post("/score_resume", response_class=JSONResponse)
async def score_resume(
    request: Request,
    resume_file: Optional[UploadFile] = File(None),
    job_file: Optional[UploadFile] = File(None),
    resume_text: Optional[str] = Form(None),
    job_description: Optional[str] = Form(None)
):
    try:
        # Get resume text from file or form
        if resume_file and resume_file.filename:
            resume_content = await process_file(resume_file)
        elif resume_text:
            resume_content = resume_text.strip()
        else:
            raise HTTPException(status_code=400, detail="Please provide a resume (file or text)")
        
        # Get job description from file or form
        if job_file and job_file.filename:
            job_content = await process_file(job_file)
        elif job_description:
            job_content = job_description.strip()
        else:
            raise HTTPException(status_code=400, detail="Please provide a job description (file or text)")
        
        if not resume_content or not job_content:
            raise HTTPException(status_code=400, detail="Resume or job description is empty")
        
        # Calculate fit score
        result = calculate_fit_score(resume_content, job_content)
        
        return JSONResponse({
            "success": True,
            "resume_text": resume_content[:500] + "..." if len(resume_content) > 500 else resume_content,
            "job_description": job_content[:500] + "..." if len(job_content) > 500 else job_content,
            **result
        })
    
    except HTTPException as e:
        return JSONResponse({"success": False, "error": e.detail}, status_code=e.status_code)
    except Exception as e:
        return JSONResponse({"success": False, "error": str(e)}, status_code=500)

@app.get("/health")
async def health_check():
    return {"status": "healthy", "model_loaded": scorer is not None}





''''from fastapi import FastAPI, Form, Request
from fastapi.responses import HTMLResponse
from fastapi.templating import Jinja2Templates
from transformers import pipeline

app = FastAPI()
templates = Jinja2Templates(directory="templates")

# Load your fine-tuned Hugging Face model
# Replace with your actual model path or use a demo model
scorer = pipeline("text-classification", model="riddhii29/results")

@app.get("/", response_class=HTMLResponse)
def home(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})

@app.post("/score_resume", response_class=HTMLResponse)
async def score_resume(request: Request,
                       resume_text: str = Form(...),
                       job_description: str = Form(...)):

    input_text = f"{resume_text} [SEP] {job_description}"
    result = scorer(input_text)[0]
    fit_score = round(float(result['score']) * 10, 1)

    return templates.TemplateResponse("index.html", {
        "request": request,
        "resume_text": resume_text,
        "job_description": job_description,
        "fit_score": fit_score
    })'''

'''from fastapi import FastAPI, File, UploadFile, Form, Request, HTTPException
from fastapi.responses import HTMLResponse, JSONResponse
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles
from transformers import pipeline
import PyPDF2
import docx
import io
from typing import Optional

app = FastAPI()
templates = Jinja2Templates(directory="templates")

# Load your fine-tuned Hugging Face model
try:
    scorer = pipeline("text-classification", model="riddhii29/results")
except Exception as e:
    print(f"Warning: Could not load model. Using demo mode. Error: {e}")
    scorer = None

def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file"""
    try:
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading PDF: {str(e)}")

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        doc_file = io.BytesIO(file_content)
        doc = docx.Document(doc_file)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text.strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading DOCX: {str(e)}")

def extract_text_from_txt(file_content: bytes) -> str:
    """Extract text from TXT file"""
    try:
        return file_content.decode('utf-8').strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading TXT: {str(e)}")

async def process_file(file: UploadFile) -> str:
    """Process uploaded file and extract text"""
    if not file:
        raise HTTPException(status_code=400, detail="No file provided")
    
    content = await file.read()
    filename = file.filename.lower()
    
    if filename.endswith('.pdf'):
        return extract_text_from_pdf(content)
    elif filename.endswith('.docx'):
        return extract_text_from_docx(content)
    elif filename.endswith('.txt'):
        return extract_text_from_txt(content)
    else:
        raise HTTPException(
            status_code=400, 
            detail="Unsupported file format. Please upload PDF, DOCX, or TXT file."
        )

def calculate_fit_score(resume_text: str, job_description: str) -> dict:
    """Calculate fit score using the model or demo scoring"""
    if scorer:
        input_text = f"{resume_text} [SEP] {job_description}"
        result = scorer(input_text)[0]
        fit_score = round(float(result['score']) * 10, 1)
        confidence = round(float(result['score']) * 100, 1)
    else:
        # Demo scoring based on keyword matching
        resume_lower = resume_text.lower()
        job_lower = job_description.lower()
        
        job_keywords = set(job_lower.split())
        resume_keywords = set(resume_lower.split())
        
        common_keywords = job_keywords.intersection(resume_keywords)
        match_ratio = len(common_keywords) / max(len(job_keywords), 1)
        
        fit_score = round(min(match_ratio * 10, 10), 1)
        confidence = round(match_ratio * 100, 1)
    
    # Determine rating category
    if fit_score >= 8:
        rating = "Excellent Match"
        color = "#10b981"
    elif fit_score >= 6:
        rating = "Good Match"
        color = "#3b82f6"
    elif fit_score >= 4:
        rating = "Fair Match"
        color = "#f59e0b"
    else:
        rating = "Poor Match"
        color = "#ef4444"
    
    return {
        "fit_score": fit_score,
        "confidence": confidence,
        "rating": rating,
        "color": color
    }

@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})

@app.post("/score_resume", response_class=JSONResponse)
async def score_resume(
    request: Request,
    resume_file: Optional[UploadFile] = File(None),
    job_file: Optional[UploadFile] = File(None),
    resume_text: Optional[str] = Form(None),
    job_description: Optional[str] = Form(None)
):
    try:
        # Get resume text from file or form
        if resume_file and resume_file.filename:
            resume_content = await process_file(resume_file)
        elif resume_text:
            resume_content = resume_text.strip()
        else:
            raise HTTPException(status_code=400, detail="Please provide a resume (file or text)")
        
        # Get job description from file or form
        if job_file and job_file.filename:
            job_content = await process_file(job_file)
        elif job_description:
            job_content = job_description.strip()
        else:
            raise HTTPException(status_code=400, detail="Please provide a job description (file or text)")
        
        if not resume_content or not job_content:
            raise HTTPException(status_code=400, detail="Resume or job description is empty")
        
        # Calculate fit score
        result = calculate_fit_score(resume_content, job_content)
        
        return JSONResponse({
            "success": True,
            "resume_text": resume_content[:500] + "..." if len(resume_content) > 500 else resume_content,
            "job_description": job_content[:500] + "..." if len(job_content) > 500 else job_content,
            **result
        })
    
    except HTTPException as e:
        return JSONResponse({"success": False, "error": e.detail}, status_code=e.status_code)
    except Exception as e:
        return JSONResponse({"success": False, "error": str(e)}, status_code=500)

@app.get("/health")
async def health_check():

    return {"status": "healthy", "model_loaded": scorer is not None}'''

